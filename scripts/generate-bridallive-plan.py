"""Generate BridalLive Integration Plan as a .docx"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUT = "/sessions/admiring-awesome-galileo/mnt/IB Daily Log/BridalLive Integration Plan.docx"

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.1)
    section.right_margin  = Inches(1.1)

# ── Helper: set paragraph font ────────────────────────────────────────────────
def set_run(run, size=11, bold=False, color=None, font="Calibri"):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_run(run, size=15, bold=True, color=(94, 56, 71))  # BridalLive plum
    # Bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '5E3847')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_run(run, size=12, bold=True, color=(60, 60, 60))
    return p

def body(text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(text)
    set_run(run, size=11)
    return p

def bullet(text, sub=False):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    if sub:
        p.paragraph_format.left_indent = Inches(0.5)
    run = p.add_run(text)
    set_run(run, size=10.5)
    return p

def status_pill(paragraph, text, bg_hex, fg=(255,255,255)):
    """Inline bold colored label"""
    run = paragraph.add_run(f"  {text}  ")
    set_run(run, size=9, bold=True, color=fg)
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), bg_hex)
    rPr.append(shd)

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.25)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(80, 40, 55)
    # Light shaded background
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F5EEF1')
    rPr.append(shd)
    return p

def add_table(headers, rows, col_widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.width = Inches(col_widths[i])
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:fill'), '5E3847')
        tcPr.append(shd)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_run(run, size=10, bold=True, color=(255,255,255))
    # Data rows
    for i, row_data in enumerate(rows):
        row = table.add_row()
        fill = 'FAF4F6' if i % 2 == 0 else 'FFFFFF'
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.width = Inches(col_widths[j])
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:fill'), fill)
            tcPr.append(shd)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_run(run, size=10)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# TITLE BLOCK
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(4)
run = p.add_run("BridalLive × IB Daily Log")
set_run(run, size=22, bold=True, color=(94, 56, 71))

p2 = doc.add_paragraph()
p2.paragraph_format.space_after = Pt(2)
run2 = p2.add_run("API Integration Plan  —  Technical Specification")
set_run(run2, size=13, color=(100, 100, 100))

p3 = doc.add_paragraph()
p3.paragraph_format.space_after = Pt(16)
run3 = p3.add_run("Prepared for Codex  ·  April 2026")
set_run(run3, size=10, color=(150, 150, 150))

# Divider
p_div = doc.add_paragraph()
pPr = p_div._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
bottom = OxmlElement('w:bottom')
bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
bottom.set(qn('w:space'), '4');    bottom.set(qn('w:color'), '5E3847')
pBdr.append(bottom); pPr.append(pBdr)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
heading1("1. Overview & Goal")
body("The IB Daily Log is a Next.js 15 / Prisma / Supabase app that tracks appointment "
     "outcomes for bridal store locations. Currently, front desk staff must manually "
     "enter every appointment into the daily log even though those appointments were "
     "already booked in BridalLive (the store's POS system). This creates duplicate "
     "data entry and makes it easy to miss entries.")
body("The goal of this integration is to pull today's scheduled appointments from "
     "BridalLive and surface them in a new Appointments section on the dashboard. "
     "Front desk can then click an appointment to check it in — pre-filling the daily "
     "log form — so they only need to record the outcome (purchased, price point, "
     "stylist time, etc.).")

heading2("Key Design Decisions")
bullet("Do NOT auto-import appointments into the daily log. Some appointments are "
       "no-shows. The check-in is always a deliberate front-desk action.")
bullet("Time Out is never pre-filled from BridalLive's scheduled end time. It must "
       "be recorded manually when the customer leaves, so stylist time-with-client "
       "metrics are accurate.")
bullet("Time In is pre-filled with BridalLive's scheduled start time as a default "
       "and can be corrected by staff.")
bullet("BridalLive is source of truth for booking data. The daily log is source of "
       "truth for outcome data.")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — CONFIRMED DETAILS
# ══════════════════════════════════════════════════════════════════════════════
heading1("2. Confirmed Account Details")

add_table(
    ["Detail", "Value"],
    [
        ["BridalLive Plan", "Elite (API access enabled)"],
        ["Account Status", "Enabled — confirmed in Settings → Account → API"],
        ["Retailer ID", "5b3d46a8"],
        ["API Key", "13192120dc53e83c  ⚠ Rotate after use"],
        ["Stores", "One BridalLive account per store location"],
        ["Auth method", "POST /apiLogin → 8-hour token"],
        ["API base URL", "https://api.bridallive.com/bl-server/api"],
    ],
    [2.2, 4.2]
)

p = doc.add_paragraph()
run = p.add_run("⚠  Note: ")
set_run(run, size=10, bold=True, color=(180, 80, 0))
run2 = p.add_run("The API auth could not be tested locally — the BridalLive server "
                 "returned 401 for all request format variations tried. This is likely "
                 "an IP allowlist restriction (calls must originate from a whitelisted "
                 "server IP). Auth should be verified once the integration is deployed "
                 "to the production server, or by contacting BridalLive support to "
                 "whitelist the development/server IP.")
set_run(run2, size=10, color=(100, 100, 100))

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — FIELD MAPPING
# ══════════════════════════════════════════════════════════════════════════════
heading1("3. Field Mapping  (BridalLive → Daily Log)")

add_table(
    ["BridalLive Field", "Daily Log Field", "Notes"],
    [
        ["First Name + Last Name",  "Customer.fullName",         "Concatenate with space"],
        ["Phone (Mobile)",          "Customer.phone",            "Already in schema, currently unused"],
        ["Email",                   "Customer.email",            "Already in schema, currently unused"],
        ["Event Date",              "appointment.wearDate",      "Direct map"],
        ["How Heard",               "leadSourceOptionId",        "Fuzzy-match to StoreOption labels"],
        ["Appointment Type",        "appointmentTypeOptionId",   "Via type-mapping config (see §5)"],
        ["Date",                    "appointmentDate",           "Direct map"],
        ["Start Time",              "timeIn",                    "Pre-filled as default; staff can correct"],
        ["End Time",                "timeOut",                   "NOT mapped — always manually entered"],
        ["Fitting Room",            "locationId",                "Match to Location by normalized name"],
        ["Associate",               "assignedStaffMemberId",     "Match to StaffMember by normalized name"],
        ["Appointment Notes",       "comments",                  "Direct map"],
        ["—",                       "purchased / pricePoint / size / status", "Staff fills in during visit"],
    ],
    [2.0, 2.2, 2.2]
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — SCHEMA CHANGES
# ══════════════════════════════════════════════════════════════════════════════
heading1("4. Schema Changes  (One Prisma Migration)")

heading2("New model: BridalLiveAppointment")
body("A staging table that holds raw synced appointments from BridalLive. "
     "Separate from the Appointment table — this represents 'what was scheduled' "
     "while Appointment represents 'what actually happened'.")

code("model BridalLiveAppointment {")
code("  id               String    @id @default(cuid())")
code("  storeId          String")
code("  store            Store     @relation(...)")
code("  bridalLiveId     String    @unique   // BridalLive's appointment ID")
code("  guestFirstName   String")
code("  guestLastName    String")
code("  guestPhone       String?")
code("  guestEmail       String?")
code("  appointmentDate  DateTime  @db.Date")
code("  scheduledStart   DateTime  // used to pre-fill timeIn")
code("  scheduledEnd     DateTime? // stored but NOT used as timeOut")
code("  appointmentType  String    // raw BridalLive label")
code("  fittingRoom      String?   // raw fitting room name")
code("  associate        String?   // raw associate name")
code("  eventDate        DateTime? @db.Date")
code("  howHeard         String?")
code("  notes            String?")
code("  isConfirmed      Boolean   @default(false)")
code("  dailyLogEntryId  String?   @unique  // set on check-in")
code("  dailyLogEntry    Appointment? @relation(...)")
code("  isNoShow         Boolean   @default(false)")
code("  syncedAt         DateTime")
code("  @@index([storeId, appointmentDate])")
code("}")

heading2("Additions to existing models")
bullet("Appointment  →  add bridalLiveAppointmentId String? @unique  (back-relation)")
bullet("Customer     →  add bridalLiveContactId String?  (for future bidirectional sync)")
bullet("Store        →  add bridalLiveSyncedAt DateTime?  (timestamp of last sync)")

heading2("New model: BridalLiveToken")
body("Stores the 8-hour auth token in the database so it survives serverless cold starts "
     "and is shared across all function instances.")
code("model BridalLiveToken {")
code("  id        String   @id  // store slug, e.g. 'galleria-curve'")
code("  token     String")
code("  expiresAt DateTime")
code("}")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — APPOINTMENT TYPE MAPPING
# ══════════════════════════════════════════════════════════════════════════════
heading1("5. Appointment Type Mapping")

body("BridalLive's appointment type labels will not match the daily log's StoreOption "
     "labels exactly. A mapping config is needed. This should live in a config file "
     "or DB table — not hardcoded — so it can be updated without a deploy.")

heading2("Visit type logic")
bullet("Walk-In bride  →  visitType: WALK_IN,  use default walk-in StoreOption type")
bullet("All other types  →  visitType: APPOINTMENT")
bullet("Alteration (Custom / 1 BM / 1 SO / 1 WG / 2 BM / 2 SO / 2 WG / 3 WG)  →  "
       "triggers the Seamstress staff field instead of Stylist in the form")

heading2("Known BridalLive types from account (confirmed via screenshot)")
add_table(
    ["BridalLive Type", "Suggested Daily Log Mapping"],
    [
        ["New Bride",              "New Bride (confirm label in StoreOptions)"],
        ["New Bride (Weekend)",    "New Bride  or  New Bride (Weekend)"],
        ["Comeback Bride",         "Comeback"],
        ["Bridesmaid",             "Bridesmaid"],
        ["Mother of Bride",        "Mother of Bride"],
        ["Destination Bride",      "Destination Bride"],
        ["Alteration (Custom)",    "Alteration"],
        ["Alteration 1/2/3 BM/SO/WG", "Alteration"],
        ["Alteration Consultation","Alteration Consultation"],
        ["Prom",                   "Prom"],
        ["Special Occasion",       "Special Occasion"],
        ["Accessories",            "Accessories"],
        ["Pick Up",                "Pick Up"],
        ["Press Pick Up",          "Press Pick Up"],
        ["Plat Press",             "Plat Press"],
        ["Interview",              "Interview"],
        ["Presentation",           "Presentation"],
        ["Customer Service",       "Customer Service"],
        ["Walk-In bride",          "→ visitType: WALK_IN + default walk-in type"],
    ],
    [2.8, 3.6]
)

body("⚠  Run SELECT label FROM StoreOption WHERE kind = 'APPOINTMENT_TYPE' in Supabase "
     "to get the exact label strings before building the mapping config.")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — FILES TO CREATE
# ══════════════════════════════════════════════════════════════════════════════
heading1("6. Files to Create")

add_table(
    ["File", "Purpose"],
    [
        ["lib/bridallive/client.ts",             "Auth + token caching (reads/writes BridalLiveToken table). Base GET/POST helpers."],
        ["lib/bridallive/types.ts",              "TypeScript types for all BridalLive API responses."],
        ["lib/bridallive/sync.ts",               "Core sync: fetch today's appointments from BridalLive API → upsert BridalLiveAppointment table."],
        ["lib/bridallive/type-mapping.ts",       "Config mapping BridalLive type labels → daily log StoreOption labels + visit type logic."],
        ["lib/server/bridallive-actions.ts",     "Server actions: syncBridalLiveAppointments(), markNoShow(), checkIn() (creates Appointment from BridalLiveAppointment)."],
        ["components/bridal-live-appointments-panel.tsx", "New dashboard section showing today's BridalLive appointments with check-in / no-show actions."],
        ["app/api/cron/sync-bridal-live/route.ts","Cron endpoint called every 30 min during business hours via Vercel Cron."],
        ["prisma/migrations/..._bridallive",     "Migration adding BridalLiveAppointment, BridalLiveToken models and new fields."],
    ],
    [2.6, 3.8]
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — DASHBOARD UX
# ══════════════════════════════════════════════════════════════════════════════
heading1("7. Dashboard UX — New Appointments Section")

body("A new panel on the dashboard shows today's BridalLive appointments. "
     "It sits above the existing check-in panel and has three row states:")

heading2("Row states")
bullet("Upcoming  —  scheduled, customer not yet checked in. "
       "Shows: guest name, scheduled time, appointment type, associate, fitting room. "
       "Actions: 'Check in' button, 'No show' button.")
bullet("Checked in  —  dailyLogEntryId is set. Green indicator. "
       "Links to the daily log entry. No further action needed.")
bullet("No show  —  isNoShow: true. Greyed out with badge. "
       "Stays visible for the day for records.")

heading2("Check-in flow")
bullet("Staff clicks 'Check in' on a BridalLiveAppointment row.")
bullet("The daily log workflow panel opens pre-filled: guest name, time in (scheduled "
       "start, editable), appointment type (mapped), stylist, location, wear date, "
       "lead source, notes.")
bullet("Time out is blank — filled in manually when customer leaves.")
bullet("Staff reviews, corrects any mismatches, and submits.")
bullet("On submit: Appointment row is created, BridalLiveAppointment.dailyLogEntryId "
       "is set. Row in dashboard turns green.")

heading2("Sync controls")
bullet("'Sync appointments' button on the dashboard — manual trigger for v1.")
bullet("Shows: 'Last synced at 2:15 PM  ·  3 upcoming  ·  1 checked in'")
bullet("If any BridalLive appointment types had no mapping match, a warning shows: "
       "'1 unrecognized type: Presentation — update type mapping in Settings.'")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — SYNC LOGIC
# ══════════════════════════════════════════════════════════════════════════════
heading1("8. Sync Logic Detail")

heading2("syncBridalLiveAppointments(storeSlug, date) algorithm")
body("Called by both the manual button and the cron job:")

p = doc.add_paragraph(style="List Number")
p.paragraph_format.space_after = Pt(3)
r = p.add_run("Get valid auth token from BridalLiveToken table. Re-auth if expired (8h TTL with 30-min buffer).")
set_run(r, size=10.5)

p = doc.add_paragraph(style="List Number")
p.paragraph_format.space_after = Pt(3)
r = p.add_run("Fetch appointments for the given date from BridalLive API.")
set_run(r, size=10.5)

p = doc.add_paragraph(style="List Number")
p.paragraph_format.space_after = Pt(3)
r = p.add_run("For each appointment: upsert into BridalLiveAppointment by bridalLiveId. "
              "Skip updating rows where dailyLogEntryId is already set (customer already checked in — don't overwrite).")
set_run(r, size=10.5)

p = doc.add_paragraph(style="List Number")
p.paragraph_format.space_after = Pt(3)
r = p.add_run("Return: { imported, updated, skipped, unmappedTypes[] }.")
set_run(r, size=10.5)

heading2("Conflict rules")
bullet("BridalLive wins for: guest name, scheduled time, appointment type, associate, fitting room.")
bullet("Daily log wins for: timeIn (actual), timeOut, purchased, pricePoint, size, status, comments.")
bullet("Once a row has dailyLogEntryId set, the sync never touches those fields again.")

heading2("No-show handling")
bullet("No-show records stay in BridalLiveAppointment with isNoShow: true.")
bullet("They do NOT create Appointment records in the daily log.")
bullet("No-show counts are queryable from BridalLiveAppointment for analytics.")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — BUILD ORDER
# ══════════════════════════════════════════════════════════════════════════════
heading1("9. Build Order")

add_table(
    ["Step", "Task", "Depends On"],
    [
        ["1", "Prisma migration — add BridalLiveAppointment, BridalLiveToken, new fields on Appointment/Customer/Store", "—"],
        ["2", "lib/bridallive/types.ts — TypeScript types (⚠ needs real API response JSON to finalize)", "BridalLive API access confirmed"],
        ["3", "lib/bridallive/client.ts — auth + token caching", "Step 1, 2"],
        ["4", "lib/bridallive/type-mapping.ts — type label mapping config", "Run StoreOption query in Supabase"],
        ["5", "lib/bridallive/sync.ts — fetch + upsert logic", "Steps 2, 3, 4"],
        ["6", "lib/server/bridallive-actions.ts — server actions (sync, markNoShow, checkIn)", "Step 5"],
        ["7", "components/bridal-live-appointments-panel.tsx — dashboard UI", "Step 6"],
        ["8", "Wire panel into app/dashboard/page.tsx", "Step 7"],
        ["9", "Connect 'Check in' button to pre-fill DailyLogWorkflowPanel", "Steps 7, 8"],
        ["10", "app/api/cron/sync-bridal-live/route.ts + vercel.json cron config", "Step 6"],
    ],
    [0.4, 4.0, 2.0]
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 10 — CURRENT STATUS & BLOCKERS
# ══════════════════════════════════════════════════════════════════════════════
heading1("10. Current Status & Blockers")

heading2("Completed")
bullet("Integration architecture fully designed and approved.")
bullet("BridalLive account confirmed: Elite plan, API enabled, Retailer ID and API Key obtained.")
bullet("All BridalLive appointment type labels confirmed via screenshot.")
bullet("Field mapping table completed.")
bullet("Schema design completed.")
bullet("explore-bridallive-api.mjs discovery script written (scripts/ directory).")

heading2("Blockers — resolve before starting Step 2")
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.left_indent  = Inches(0.25)
run = p.add_run("BLOCKER 1  ")
set_run(run, size=10.5, bold=True, color=(180, 50, 50))
run2 = p.add_run("API auth returns 401 when called from local machine. "
                 "Likely an IP allowlist restriction on BridalLive's server. "
                 "Resolution: contact BridalLive support and ask them to whitelist "
                 "the production server IP (and optionally dev IP). "
                 "Alternatively, deploy the integration to production first and test there.")
set_run(run2, size=10.5)

p2 = doc.add_paragraph()
p2.paragraph_format.space_before = Pt(4)
p2.paragraph_format.left_indent  = Inches(0.25)
run3 = p2.add_run("BLOCKER 2  ")
set_run(run3, size=10.5, bold=True, color=(180, 100, 0))
run4 = p2.add_run("Exact BridalLive API response JSON shapes are unknown (needed for types.ts). "
                  "Resolution: once auth is working, run scripts/explore-bridallive-api.mjs "
                  "and paste the output bridallive-api-shapes.json. Steps 2-5 can then be "
                  "completed in one pass.")
set_run(run4, size=10.5)

p3 = doc.add_paragraph()
p3.paragraph_format.space_before = Pt(4)
p3.paragraph_format.left_indent  = Inches(0.25)
run5 = p3.add_run("BLOCKER 3  ")
set_run(run5, size=10.5, bold=True, color=(100, 100, 0))
run6 = p3.add_run("StoreOption appointment type labels unknown. "
                  "Resolution: run the following query in Supabase SQL editor and "
                  "share the output so type-mapping.ts can be built accurately:")
set_run(run6, size=10.5)
code('SELECT label FROM "StoreOption" WHERE kind = \'APPOINTMENT_TYPE\' ORDER BY label;')

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 11 — ENV VARS
# ══════════════════════════════════════════════════════════════════════════════
heading1("11. Environment Variables to Add")

add_table(
    ["Variable", "Value", "Where"],
    [
        ["BRIDALLIVE_RETAILER_ID", "5b3d46a8", ".env.local + Vercel/Netlify dashboard"],
        ["BRIDALLIVE_API_KEY",     "13192120dc53e83c  (rotate first)", ".env.local + Vercel/Netlify dashboard"],
        ["CRON_SECRET",            "Generate a random 32-char string", ".env.local + deployment dashboard"],
    ],
    [2.4, 2.8, 1.4]
)

body("The CRON_SECRET is used to protect the /api/cron/sync-bridal-live endpoint so "
     "only the cron scheduler can trigger it.")

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
doc.save(OUT)
print(f"Saved: {OUT}")
